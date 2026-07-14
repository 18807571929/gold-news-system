"""将毕设 Markdown 章节导出为 Word（.docx），套用学校论文模板样式。

⚠️ 安全提示：已插图/手工编辑的正文请用 patch_thesis_docx.py，不要用本脚本覆盖。

用法:
  # 推荐：从 Markdown 重建到新文件（不碰正文定稿）
  python scripts/generate_thesis_docx.py --chapters 1 2 3 4 5 6 7

  # 只导出部分章节
  python scripts/generate_thesis_docx.py --chapters 4 5

  # 确需覆盖已有文件（含 毕设正文_全文.docx）时必须加 --force
  python scripts/generate_thesis_docx.py --out docs/毕设/毕设正文_全文.docx --force

日常更新已编辑 Word → 请用: python scripts/patch_thesis_docx.py
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

THESIS_DIR = PROJECT_ROOT / "docs" / "毕设"

# 禁止在无 --force 时写入的文件名（已定稿、已插图版本）
PROTECTED_OUTPUT_NAMES = frozenset({
    "毕设正文_全文.docx",
})

# 默认重建输出（与定稿文件名区分，避免误覆盖）
DEFAULT_REBUILD_OUTPUT = THESIS_DIR / "毕设正文_从Markdown重建.docx"

# 学校模板（.doc 需 Windows + Word 首次转换为同目录 .docx 缓存）
TEMPLATE_DOC = Path(r"E:\HomeWork\朱启亮毕业论文\2026届（2022级）-论文模板-智能系统方向.doc")
TEMPLATE_DOCX = THESIS_DIR / "_template_智能系统方向.docx"

# 章号 -> Markdown 文件名（相对 docs/毕设/）
# 值为 None 表示尚未撰写，导出时自动跳过并提示
CHAPTER_FILES: dict[int, str | None] = {
    1: "01_第1章_绪论.md",
    2: "02_第2章_相关技术与理论.md",
    3: "03_第3章_系统总体设计.md",
    4: "04_第4章_新闻冲击网格适配策略.md",
    5: "05_第5章_回测与实验.md",
    6: "06_第6章_MT5模拟盘验证.md",
    7: "07_第7章_总结与展望.md",
}

THESIS_TITLE = "黄金新闻实时利多利空分析与趋势强度判断系统"
THESIS_SUBTITLE = "（智能系统方向 · 毕设正文）"
AUTHOR = "朱启亮"
SCHOOL = "珠海科技学院"


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit("请先安装: pip install python-docx") from exc
    return Document, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt, RGBColor


def ensure_template_docx() -> Path | None:
    """将学校 .doc 模板转为 .docx 缓存；无 Word 时返回 None。"""
    if not TEMPLATE_DOC.exists():
        print(f"警告: 未找到模板 {TEMPLATE_DOC}，将使用脚本内置样式")
        return None

    if TEMPLATE_DOCX.exists() and TEMPLATE_DOCX.stat().st_mtime >= TEMPLATE_DOC.stat().st_mtime:
        return TEMPLATE_DOCX

    try:
        import win32com.client
    except ImportError:
        print("警告: 未安装 pywin32，无法转换 .doc 模板，将使用脚本内置样式")
        return TEMPLATE_DOCX if TEMPLATE_DOCX.exists() else None

    TEMPLATE_DOCX.parent.mkdir(parents=True, exist_ok=True)
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(TEMPLATE_DOC.resolve()))
        doc.SaveAs2(str(TEMPLATE_DOCX.resolve()), FileFormat=16)
        doc.Close()
        print(f"已转换模板: {TEMPLATE_DOCX}")
        return TEMPLATE_DOCX
    finally:
        word.Quit()


def clear_document_body(doc) -> None:
    """清空正文段落与表格，保留样式与节属性。"""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag != "sectPr":
            body.remove(child)


def setup_thesis_styles(doc) -> None:
    """按学校模板：A4、2.5cm 页边距、宋体小四正文、黑体标题、1.25 倍行距。"""
    _, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt, _ = _require_docx()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    def _set_style(
        style_name: str,
        font_name: str,
        size_pt: float,
        *,
        bold: bool | None = None,
        align=None,
        line_spacing: float | None = 1.25,
        first_indent_cm: float | None = None,
        space_before_pt: float = 0,
        space_after_pt: float = 0,
    ) -> None:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size_pt)
        if bold is not None:
            style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        pf = style.paragraph_format
        if line_spacing is not None:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = line_spacing
        else:
            pf.line_spacing_rule = None
            pf.line_spacing = None
        if align is not None:
            pf.alignment = align
        pf.first_line_indent = Cm(first_indent_cm) if first_indent_cm else Pt(0)
        pf.space_before = Pt(space_before_pt)
        pf.space_after = Pt(space_after_pt)

    _set_style("Normal", "宋体", 12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent_cm=0.74)
    _set_style("Heading 1", "黑体", 22, bold=True, line_spacing=2.4, space_before_pt=17, space_after_pt=16.5)
    _set_style("Heading 2", "黑体", 16, bold=True, line_spacing=1.73, space_before_pt=13, space_after_pt=13)
    _set_style("Heading 3", "黑体", 14, bold=True, space_before_pt=13, space_after_pt=13)

    try:
        _set_style("Caption", "宋体", 10.5, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=None, first_indent_cm=None)
    except KeyError:
        pass


def open_base_document():
    """优先基于学校模板 docx（仅保留样式），否则新建文档。"""
    Document, _, _, _, _, _, _ = _require_docx()
    template_path = ensure_template_docx()
    if template_path and template_path.exists():
        doc = Document(str(template_path))
        clear_document_body(doc)
        if len(doc.sections) > 1:
            while len(doc.sections) > 1:
                doc.sections[-1]._sectPr.getparent().remove(doc.sections[-1]._sectPr)
        setup_thesis_styles(doc)
        return doc
    doc = Document()
    setup_thesis_styles(doc)
    return doc


def _set_run_font(run, qn, font_name: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    from docx.shared import Pt

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_cover_page(doc) -> None:
    """封面：对齐学校模板居中黑体标题 + 元信息。"""
    _, WD_ALIGN_PARAGRAPH, _, qn, _, Pt, _ = _require_docx()

    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(12)
    r0 = p0.add_run(SCHOOL)
    _set_run_font(r0, qn, "宋体", 26)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(24)
    r1 = p1.add_run("毕 业 论 文")
    _set_run_font(r1, qn, "黑体", 26, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(36)
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(THESIS_TITLE)
    _set_run_font(r2, qn, "黑体", 18, bold=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(THESIS_SUBTITLE)
    _set_run_font(r3, qn, "宋体", 14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(72)
    for label, val in [("作者：", AUTHOR), ("日期：", "2026年7月12日"), ("项目：", str(PROJECT_ROOT))]:
        run = meta.add_run(f"{label}{val}\n")
        _set_run_font(run, qn, "宋体", 12)

    doc.add_page_break()


def add_body_paragraph(doc, text: str, *, indent: bool = True, bold: bool = False) -> None:
    _, _, WD_LINE_SPACING, qn, _, Pt, _ = _require_docx()
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25
    pf.first_line_indent = Pt(24) if indent and text.strip() else Pt(0)
    pf.space_after = Pt(6)
    run = para.add_run(text)
    _set_run_font(run, qn, "宋体", 12, bold=bold)


def add_table_caption(doc, text: str) -> None:
    """表题：居中，如「表4-1 …」。"""
    _, WD_ALIGN_PARAGRAPH, _, qn, _, Pt, _ = _require_docx()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    _set_run_font(run, qn, "宋体", 12, bold=True)


def add_figure_placeholder(doc, fig_id: str, caption: str) -> None:
    """图占位：居中虚线框 + 图题（如 图4-1 方向解析流程图）。"""
    _, WD_ALIGN_PARAGRAPH, _, qn, Cm, Pt, RGBColor = _require_docx()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as oxml_qn

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(f"【此处插入{caption}】")
    _set_run_font(run, qn, "宋体", 12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(oxml_qn("w:tcBorders")):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(oxml_qn("w:val"), "dashed")
        el.set(oxml_qn("w:sz"), "8")
        el.set(oxml_qn("w:color"), "999999")
        borders.append(el)
    tc_pr.append(borders)
    for old in tc_pr.findall(oxml_qn("w:tcW")):
        tc_pr.remove(old)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(oxml_qn("w:type"), "pct")
    tc_w.set(oxml_qn("w:w"), "8000")
    tc_pr.append(tc_w)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    cap_run = cap.add_run(f"图{fig_id} {caption}")
    _set_run_font(cap_run, qn, "宋体", 10.5)

    doc.add_paragraph("")


def add_blockquote(doc, text: str) -> None:
    _, _, WD_LINE_SPACING, qn, Cm, Pt, RGBColor = _require_docx()
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25
    pf.left_indent = Cm(0.74)
    pf.space_after = Pt(6)
    run = para.add_run(text)
    _set_run_font(run, qn, "楷体", 12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True


def add_code_block(doc, lines: list[str]) -> None:
    _, _, WD_LINE_SPACING, qn, Cm, Pt, RGBColor = _require_docx()
    text = "\n".join(lines)
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(0.5)
    pf.space_after = Pt(6)
    run = para.add_run(text)
    _set_run_font(run, qn, "Consolas", 10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def add_markdown_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    _, _, _, qn, _, _, _ = _require_docx()
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, qn, "宋体", 11, bold=(ri == 0))
    doc.add_paragraph("")


def heading_level_from_text(text: str) -> int | None:
    m = re.match(r"^(#{1,6})\s+(.+)$", text.strip())
    if not m:
        return None
    return min(len(m.group(1)), 3)


def heading_text(text: str) -> str:
    return re.sub(r"^#{1,6}\s+", "", text.strip())


def parse_inline_bold(paragraph, text: str, qn) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, qn, "宋体", 12, bold=True)
        elif part:
            run = paragraph.add_run(part)
            _set_run_font(run, qn, "宋体", 12)


def is_table_caption_line(line: str) -> bool:
    s = line.strip()
    if s.startswith("**") and s.endswith("**"):
        inner = s[2:-2].strip()
        return inner.startswith("表")
    return False


def extract_table_caption(line: str) -> str:
    s = line.strip()
    if s.startswith("**") and s.endswith("**"):
        return s[2:-2].strip()
    return s.strip()


def markdown_to_docx(doc, md_path: Path) -> None:
    _, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, _, Pt, _ = _require_docx()
    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            add_markdown_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if in_code:
            if line.strip().startswith("```"):
                if code_lang == "mermaid":
                    add_figure_placeholder(doc, "4-1", "方向解析流程图")
                else:
                    add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if is_table_separator(line):
                i += 1
                continue
            table_buf.append(parse_table_row(line))
            i += 1
            continue
        flush_table()

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith(">"):
            quote = re.sub(r"^>\s*", "", line.strip())
            add_blockquote(doc, quote)
            i += 1
            continue

        if is_table_caption_line(line):
            add_table_caption(doc, extract_table_caption(line))
            i += 1
            continue

        hl = heading_level_from_text(line)
        if hl is not None:
            text = heading_text(line)
            if text.startswith("第") and "章" in text[:8]:
                h = doc.add_heading(text, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in h.runs:
                    _set_run_font(run, qn, "黑体", 22, bold=True)
            else:
                h = doc.add_heading(text, level=hl)
                for run in h.runs:
                    _set_run_font(run, qn, "黑体", {1: 22, 2: 16, 3: 14}.get(hl, 14), bold=True)
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m:
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.25
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(24)
            pf.space_after = Pt(6)
            num_run = para.add_run(f"{m.group(1)}. ")
            _set_run_font(num_run, qn, "宋体", 12)
            parse_inline_bold(para, m.group(2), qn)
            i += 1
            continue

        if line.strip().startswith("- "):
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.25
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(24)
            pf.space_after = Pt(6)
            bullet_run = para.add_run("• ")
            _set_run_font(bullet_run, qn, "宋体", 12)
            parse_inline_bold(para, line.strip()[2:], qn)
            i += 1
            continue

        if line.strip().startswith("**") and line.strip().endswith("**") and line.count("**") == 2:
            add_body_paragraph(doc, line.strip()[2:-2], indent=True, bold=True)
            i += 1
            continue

        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.25
        pf.first_line_indent = Pt(24)
        pf.space_after = Pt(6)
        parse_inline_bold(para, line.strip(), qn)
        i += 1

    flush_table()
    if in_code and code_lines:
        if code_lang == "mermaid":
            add_figure_placeholder(doc, "4-1", "方向解析流程图")
        else:
            add_code_block(doc, code_lines)


def build_thesis_docx(chapters: list[int], out_path: Path) -> None:
    doc = open_base_document()
    add_cover_page(doc)

    written = 0
    for idx, ch in enumerate(chapters):
        md_name = CHAPTER_FILES.get(ch)
        if md_name is None:
            print(f"跳过第{ch}章: CHAPTER_FILES[{ch}] 为 None，请在脚本中填入 md 文件名")
            continue
        if not md_name:
            print(f"跳过未知章节: {ch}")
            continue
        md_path = THESIS_DIR / md_name
        if not md_path.exists():
            print(f"文件不存在: {md_path}")
            continue
        print(f"写入: 第{ch}章 <- {md_path.name}")
        markdown_to_docx(doc, md_path)
        written += 1
        if idx < len(chapters) - 1:
            doc.add_page_break()

    if written == 0:
        raise SystemExit("未写入任何章节，请检查 --chapters 与 CHAPTER_FILES")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已生成: {out_path}")


def assert_safe_output_path(out_path: Path, force: bool) -> None:
    """防止误覆盖已定稿 Word（含用户已插图表）。"""
    resolved = out_path.resolve()
    name = resolved.name

    if name in PROTECTED_OUTPUT_NAMES and not force:
        raise SystemExit(
            f"\n❌ 拒绝写入受保护文件: {resolved}\n\n"
            "该文件通常已手动插入图 3-1 / 图 4-1 等，全量重建会丢失这些内容。\n\n"
            "请选择:\n"
            "  1) 从 Markdown 重建到新文件（推荐）:\n"
            f"     python scripts/generate_thesis_docx.py --chapters 1 2 3 4 5 6 7\n"
            f"     → 默认输出 {DEFAULT_REBUILD_OUTPUT}\n\n"
            "  2) 只更新图/参考文献，保留正文:\n"
            "     python scripts/patch_thesis_docx.py\n\n"
            "  3) 确知要覆盖定稿，自行承担丢失插图风险:\n"
            f"     python scripts/generate_thesis_docx.py --out \"{resolved}\" --force\n"
        )

    if resolved.exists() and not force:
        raise SystemExit(
            f"\n❌ 输出文件已存在: {resolved}\n\n"
            "为避免误覆盖，已中止。可改用新文件名，例如:\n"
            f"  --out \"{DEFAULT_REBUILD_OUTPUT}\"\n\n"
            "或确认要覆盖后加 --force。\n"
            "日常改已定稿请用: python scripts/patch_thesis_docx.py\n"
        )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="毕设 Markdown → Word（学校模板）。默认写入「从Markdown重建」新文件，不覆盖定稿。",
    )
    parser.add_argument(
        "--chapters",
        type=int,
        nargs="+",
        default=[1, 2, 3, 4, 5, 6, 7],
        help="要导出的章节号（默认 1–7）",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=DEFAULT_REBUILD_OUTPUT,
        help=f"输出 .docx 路径（默认 {DEFAULT_REBUILD_OUTPUT.name}）",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="允许覆盖已存在的输出文件（含 毕设正文_全文.docx）",
    )
    args = parser.parse_args()

    out_path = args.out
    if not out_path.is_absolute():
        out_path = (PROJECT_ROOT / out_path).resolve()

    assert_safe_output_path(out_path, args.force)

    if args.force:
        print("⚠️  --force 已启用：将覆盖 ", out_path)

    build_thesis_docx(args.chapters, out_path)


if __name__ == "__main__":
    main()
