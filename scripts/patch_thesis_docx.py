"""在不重建全文的前提下，向已有毕设 Word 插入图3-1、更新参考文献。

✅ 日常更新已定稿 Word 请优先用本脚本（保留图 3-1 / 图 4-1 等手工内容）。

从 Markdown 全量重建 → 请用 generate_thesis_docx.py（默认写入新文件，不覆盖定稿）

用法:
  python scripts/patch_thesis_docx.py
  python scripts/patch_thesis_docx.py --docx "E:\\HomeWork\\朱启亮毕业论文\\毕设正文_全文.docx"
"""

from __future__ import annotations

import argparse
import shutil
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG3_PNG = PROJECT_ROOT / "docs" / "毕设" / "assets" / "fig3-1_系统架构.png"

REFERENCES = [
    "[1] Tetlock P C. Giving content to investor sentiment: The role of media in the stock market[J]. "
    "The Journal of Finance, 2007, 62(3): 1139-1168.",
    "[2] Loughran T, McDonald B. Textual analysis in accounting and finance: A survey[J]. "
    "Journal of Accounting Research, 2016, 54(4): 1187-1230.",
    "[3] Kearns J, Miskolczi L. Macro news and commodity returns[R/OL]. Federal Reserve Board, 2021.",
    "[4] MetaQuotes Ltd. MetaTrader 5 Python API Documentation[EB/OL]. "
    "https://www.mql5.com/en/docs/python_metatrader5, 2024.",
    "[5] Federal Reserve Bank of St. Louis. FRED Economic Data[EB/OL]. "
    "https://fred.stlouisfed.org/, 2024.",
    "[6] 姚加权, 冯绪, 王赞, 等. 语调、情绪及市场影响：基于金融情绪词典[J]. "
    "管理科学学报, 2020, 23(5): 26-46.",
    "[7] 张维, 张永杰, 周洪荣. 计算实验金融研究进展[J]. 管理科学学报, 2010, 13(6): 1-11.",
    "[8] 珠海科技学院. 2026届（2022级）毕业论文模板（智能系统方向）[Z]. 2025.",
]


def _require_docx():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    return Document, WD_ALIGN_PARAGRAPH, qn, Cm, Pt


def _set_run_font(run, qn, font_name: str, size_pt: float) -> None:
    from docx.shared import Pt

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)


def backup_docx(path: Path) -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = path.with_name(f"{path.stem}_patch前备份_{stamp}{path.suffix}")
    shutil.copy2(path, bak)
    print(f"已备份: {bak}")
    return bak


def insert_fig3(doc, png: Path) -> bool:
    Document, WD_ALIGN_PARAGRAPH, qn, Cm, Pt = _require_docx()
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    if not png.exists():
        print(f"警告: 未找到 {png}，跳过图3-1")
        return False

    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if "图3-1" in t and ("架构" in t or "示意" in t):
            target_idx = i
            break

    if target_idx is None:
        print("警告: 未定位图3-1插入位置")
        return False

    placeholder = doc.paragraphs[target_idx]
    placeholder.text = ""
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = placeholder.add_run()
    run.add_picture(str(png), width=Cm(15))

    new_p = OxmlElement("w:p")
    placeholder._p.addnext(new_p)
    cap = Paragraph(new_p, placeholder._parent)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    cap_run = cap.add_run("图3-1 系统总体架构示意图")
    _set_run_font(cap_run, qn, "宋体", 10.5)

    print(f"已插入图3-1 @ 段落 {target_idx}")
    return True


def update_references(doc) -> bool:
    _, WD_ALIGN_PARAGRAPH, qn, _, Pt = _require_docx()
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    ref_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "参考文献" in para.text:
            ref_idx = i
            break

    if ref_idx is None:
        print("警告: 未找到参考文献标题")
        return False

    to_remove = []
    for j in range(ref_idx + 1, len(doc.paragraphs)):
        t = doc.paragraphs[j].text.strip()
        if not t:
            to_remove.append(doc.paragraphs[j])
            continue
        if t.startswith("[") or (t[0].isdigit() and "." in t[:4]):
            to_remove.append(doc.paragraphs[j])
        else:
            break

    for para in to_remove:
        p = para._element
        p.getparent().remove(p)

    title = doc.paragraphs[ref_idx]
    title.text = "参考文献"
    for run in title.runs:
        _set_run_font(run, qn, "黑体", 14)

    anchor = title
    for ref in REFERENCES:
        new_p = OxmlElement("w:p")
        anchor._p.addnext(new_p)
        para = Paragraph(new_p, anchor._parent)
        para.paragraph_format.line_spacing = 1.25
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(ref)
        _set_run_font(run, qn, "宋体", 10.5)
        anchor = para

    print(f"已更新参考文献 {len(REFERENCES)} 条")
    return True


def patch(docx_path: Path, fig_png: Path) -> None:
    Document, _, _, _, _ = _require_docx()
    backup_docx(docx_path)
    doc = Document(str(docx_path))
    insert_fig3(doc, fig_png)
    update_references(doc)
    doc.save(str(docx_path))
    print(f"已保存: {docx_path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--docx",
        type=Path,
        default=Path(r"E:\HomeWork\朱启亮毕业论文\毕设正文_全文.docx"),
    )
    parser.add_argument("--fig", type=Path, default=FIG3_PNG)
    args = parser.parse_args()
    if not args.docx.exists():
        raise SystemExit(f"文件不存在: {args.docx}")
    patch(args.docx, args.fig)
    # 同步项目副本
    proj_copy = PROJECT_ROOT / "docs" / "毕设" / "毕设正文_全文.docx"
    if args.docx.resolve() != proj_copy.resolve():
        shutil.copy2(args.docx, proj_copy)
        print(f"已同步: {proj_copy}")


if __name__ == "__main__":
    main()
