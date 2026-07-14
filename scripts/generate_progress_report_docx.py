"""生成 Word 版进度汇报（给老师）。"""

from __future__ import annotations

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))


def build_docx(out_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm
    except ImportError:
        raise SystemExit("请先安装: pip install python-docx")

    doc = Document()

    # 中文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = doc.add_heading("黄金新闻实时利多利空分析与趋势强度判断系统", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("阶段性进度汇报")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.add_run("汇报人：").bold = True
    info.add_run("朱启亮\n")
    info.add_run("汇报日期：").bold = True
    info.add_run("2026年7月11日\n")
    info.add_run("项目路径：").bold = True
    info.add_run("E:\\gold-news-system\n")
    info.add_run("当前版本：").bold = True
    info.add_run("Phase 0–4 主体开发完成")

    doc.add_heading("一、项目概述", level=1)
    doc.add_paragraph(
        "本课题构建黄金（GOLD）新闻实时分析系统，实现多源新闻抓取、LLM/词典情感分析、"
        "五因子趋势强度评分（L1–L4）、网格策略适配及 MT5 模拟盘执行，并与量化回测框架对接，"
        "验证新闻冲击下网格策略的防护效果。"
    )
    doc.add_paragraph("技术栈：Python 3.10+、DeepSeek API、MetaTrader5、Parquet 归档。")

    doc.add_heading("二、总体进度", level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    hdr = ["阶段", "名称", "状态", "完成度"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ("Phase 0", "环境搭建与数据管道", "已完成", "100%"),
        ("Phase 1", "NLP 情感分析引擎", "已完成", "100%"),
        ("Phase 2", "趋势评分与网格策略", "已完成", "100%"),
        ("Phase 3", "MT5 模拟盘集成", "代码完成", "95%"),
        ("Phase 4", "回测验证 A/B", "已完成", "100%"),
    ]
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph(
        "说明：Phase 3 的 L2 级删反向挂单实盘验证因周末市场休市（retcode 10018）"
        "待下一交易时段执行，其余代码与连接测试已通过。"
    )

    doc.add_heading("三、各阶段主要成果", level=1)

    doc.add_heading("3.1 Phase 0 — 数据管道", level=2)
    for item in [
        "JSON 缓存同步 Parquet 至量化项目 History Data",
        "MT5 环境诊断脚本 diagnose_mt5_env.py",
        "路径配置支持 GOLD_DATA_ROOT 环境变量",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Phase 1 — NLP 引擎", level=2)
    for item in [
        "金十数据 + 东方财富双源抓取",
        "DeepSeek API 四跳 Chain-of-Thought，支持多事件解析",
        "词典扩充与多源共识因子",
        "test_nlp_api.py 验证通过",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.3 Phase 2 — 趋势与网格", level=2)
    for item in [
        "五因子权重：0.35/0.15/0.20/0.15/0.15",
        "L1–L4 趋势等级与网格动作规则对齐设计文档",
        "ATR 动态网格间距（MT5 实时 + CSV 兜底）",
        "事件持续性 duration_estimator",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.4 Phase 3 — MT5 集成", level=2)
    for item in [
        "FxPro 模拟盘 591838672 @ FxPro-MT5 Demo，品种 GOLD",
        "connector.py 连接验证：connected=true，bid≈4120",
        "grid_executor：删反向挂单、部分平仓、L2 顺势布网格",
        "risk_state 风控状态机：常规/预警/紧急",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.5 Phase 4 — 回测验证", level=2)
    doc.add_paragraph(
        "正式回测窗口：2026-06-27 ~ 2026-07-10（H1，14天），30条真实新闻信号。"
    )

    bt = doc.add_table(rows=4, cols=4)
    bt.style = "Table Grid"
    bt.rows[0].cells[0].text = "指标"
    bt.rows[0].cells[1].text = "Baseline"
    bt.rows[0].cells[2].text = "News Adapter"
    bt.rows[0].cells[3].text = "差值"
    bt_data = [
        ("总盈亏", "+346", "+526", "+180"),
        ("最大回撤", "13.91%", "13.91%", "0"),
        ("冲击动作", "0", "14", "—"),
    ]
    for i, row in enumerate(bt_data, 1):
        for j, v in enumerate(row):
            bt.rows[i].cells[j].text = v

    doc.add_paragraph(
        "结论：新闻冲击适配组较 baseline 多盈利 180 美元（初始 10000），"
        "最大回撤持平，初步验证防护策略有效。"
    )

    doc.add_heading("四、系统数据流", level=1)
    doc.add_paragraph(
        "新闻抓取 → NLP 情感分析 → 五因子趋势评分 → GridAdapter 网格建议 "
        "→ SignalBridge 信号 JSON → GridExecutor MT5 执行；"
        "并行归档至 Parquet 并供 backtest_adapter 做 A/B 回测。"
    )

    doc.add_heading("五、待完成事项", level=1)
    for item in [
        "L2 实盘验证（删反向挂单）+ MT5 截图（需交易时段）",
        "持续运行积累更多信号样本（建议 1 周）",
        "扩展回测窗口至 30~90 天",
        "与 gs_v17 全量网格引擎深度集成",
        "毕设正文系统设计与实验章节撰写",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("六、主要文件位置", level=1)
    paths = [
        ("源代码", "E:\\gold-news-system\\src\\"),
        ("操作记录", "E:\\gold-news-system\\docs\\操作记录\\"),
        ("回测报告", "E:\\量化项目\\05 Backtest\\results\\news-grid\\"),
        ("框架规划", "E:\\量化项目\\01 thesis\\黄金新闻分析系统_框架规划.md"),
    ]
    for name, path in paths:
        p = doc.add_paragraph()
        p.add_run(f"{name}：").bold = True
        p.add_run(path)

    doc.add_paragraph("")
    footer = doc.add_paragraph("汇报人：朱启亮    日期：2026年7月11日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已生成: {out_path}")


if __name__ == "__main__":
    md_dir = PROJECT_ROOT / "docs" / "汇报"
    out = md_dir / "黄金新闻分析系统_进度汇报.docx"
    build_docx(out)
